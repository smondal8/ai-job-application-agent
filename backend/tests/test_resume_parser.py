import io
import json
import docx
import pypdf
import pytest
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.core.database import Base
from app.core.errors import BadRequestError
from app.models.candidate import CandidateProfile
from app.services.profile_service import CandidateProfileService
from app.services.resume_parser_service import ResumeParserService


# --- Realistic RxResume Official Schema Fixture ---
RX_RESUME_ACTUAL_FIXTURE = """
{
  "basics": {
    "name": "Dr. Elena Rostova",
    "headline": "Principal AI Research Scientist & Quantum Engineer",
    "email": "elena.rostova@quantum-labs.io",
    "phone": "+1 (415) 888-9900",
    "location": "San Francisco, CA",
    "website": {
      "url": "https://elenarostova.ai",
      "label": "Personal Website"
    },
    "customFields": [
      {
        "id": "cf-1",
        "icon": "globe",
        "text": "Security Clearance: Top Secret",
        "link": ""
      }
    ]
  },
  "summary": {
    "title": "Professional Summary",
    "content": "<p>Lead AI research scientist specializing in quantum-enhanced representation learning and distributed LLM alignment.</p>",
    "hidden": false
  },
  "sections": {
    "profiles": {
      "title": "Profiles",
      "hidden": false,
      "items": [
        {
          "id": "prof-1",
          "hidden": false,
          "icon": "linkedin-logo",
          "network": "LinkedIn",
          "username": "elena-rostova",
          "website": {
            "url": "https://linkedin.com/in/elena-rostova",
            "label": "LinkedIn Profile"
          }
        },
        {
          "id": "prof-2",
          "hidden": false,
          "icon": "github-logo",
          "network": "GitHub",
          "username": "erostova",
          "website": {
            "url": "https://github.com/erostova",
            "label": "GitHub Repositories"
          }
        }
      ]
    },
    "experience": {
      "title": "Experience",
      "hidden": false,
      "items": [
        {
          "id": "exp-1",
          "hidden": false,
          "company": "Quantum AI Dynamics",
          "position": "Staff AI Research Scientist",
          "location": "San Francisco, CA",
          "period": "2021 - Present",
          "website": {
            "url": "https://quantum-ai.example",
            "label": "Company Website"
          },
          "description": "<ul><li>Architected fault-tolerant tensor contraction engines.</li><li>Published 4 peer-reviewed papers on LLM verification.</li></ul>"
        }
      ]
    },
    "education": {
      "title": "Education",
      "hidden": false,
      "items": [
        {
          "id": "edu-1",
          "hidden": false,
          "school": "Stanford University",
          "degree": "Ph.D.",
          "area": "Computer Science & Quantum Computing",
          "grade": "4.0",
          "location": "Stanford, CA",
          "period": "2016 - 2021",
          "description": "Dissertation on quantum neural networks."
        }
      ]
    },
    "skills": {
      "title": "Skills",
      "hidden": false,
      "items": [
        {
          "id": "sk-1",
          "hidden": false,
          "icon": "code",
          "name": "PyTorch",
          "proficiency": "Expert",
          "level": 5,
          "keywords": ["Distributed Training", "CUDA", "TensorRT"]
        },
        {
          "id": "sk-2",
          "hidden": false,
          "icon": "code",
          "name": "Rust",
          "proficiency": "Advanced",
          "level": 4,
          "keywords": ["Tokio", "Rayon"]
        }
      ]
    },
    "projects": {
      "title": "Projects",
      "hidden": false,
      "items": [
        {
          "id": "proj-1",
          "hidden": false,
          "name": "Quantum-Qwen Engine",
          "period": "2023",
          "website": {
            "url": "https://github.com/erostova/quantum-qwen",
            "label": "Project Link"
          },
          "description": "<p>Accelerated hybrid quantum-classical attention kernel.</p>"
        }
      ]
    }
  }
}
"""


def test_parse_rxresume_actual_schema_profile_and_all_sections():
    """Test 100% mapping of official RxResume JSON schema into Candidate Profile, Experience, Education, Skills, and Projects."""
    parser = ResumeParserService()
    clean_text, result = parser.parse_file_bytes(
        RX_RESUME_ACTUAL_FIXTURE.encode("utf-8"),
        filename="rxresume_export.json",
        mime_type="application/json",
    )

    # Invariant checks
    assert result["provenance"] == "untrusted_import"
    assert result["is_verified"] is False

    profile = result["profile"]

    # 1. Assert full_name
    assert profile["full_name"] == "Dr. Elena Rostova"

    # 2. Assert email
    assert profile["email"] == "elena.rostova@quantum-labs.io"

    # 3. Assert phone
    assert profile["phone"] == "+1 (415) 888-9900"

    # 4. Assert location
    assert profile["location"] == "San Francisco, CA"

    # 5. Assert headline
    assert profile["headline"] == "Principal AI Research Scientist & Quantum Engineer"

    # 6. Assert summary (HTML tags stripped)
    assert profile["summary"] == "Lead AI research scientist specializing in quantum-enhanced representation learning and distributed LLM alignment."

    # 7. Assert LinkedIn URL
    assert profile["linkedin_url"] == "https://linkedin.com/in/elena-rostova"

    # 8. Assert GitHub URL
    assert profile["github_url"] == "https://github.com/erostova"

    # Assert website / portfolio URL
    assert profile["portfolio_url"] == "https://elenarostova.ai"

    # 9. Assert experience extraction remains intact
    assert len(result["experiences"]) == 1
    exp = result["experiences"][0]
    assert exp["company"] == "Quantum AI Dynamics"
    assert exp["position"] == "Staff AI Research Scientist"
    assert exp["location"] == "San Francisco, CA"
    assert exp["is_current"] is True
    assert len(exp["highlights"]) == 2
    assert "Architected fault-tolerant tensor contraction engines." in exp["highlights"]

    # 10. Assert education, skills, and projects remain intact
    assert len(result["educations"]) == 1
    edu = result["educations"][0]
    assert edu["institution"] == "Stanford University"
    assert edu["degree"] == "Ph.D."
    assert edu["field_of_study"] == "Computer Science & Quantum Computing"
    assert edu["gpa"] == "4.0"

    skill_names = [s["name"] for s in result["skills"]]
    assert "PyTorch" in skill_names
    assert "Rust" in skill_names
    assert "Distributed Training" in skill_names
    assert "CUDA" in skill_names

    assert len(result["projects"]) == 1
    proj = result["projects"][0]
    assert proj["name"] == "Quantum-Qwen Engine"
    assert proj["url"] == "https://github.com/erostova/quantum-qwen"
    assert "Accelerated hybrid quantum-classical attention kernel." in proj["description"]


def test_apply_rxresume_import_to_database_candidate_profile():
    """Verify that importing and applying RxResume JSON populates all candidate profile columns in SQLite."""
    engine = create_engine("sqlite:///:memory:")
    Base.metadata.create_all(bind=engine)
    SessionLocal = sessionmaker(bind=engine)
    db = SessionLocal()

    profile_service = CandidateProfileService()
    master_profile = profile_service.get_or_create_primary_profile(db)

    raw_import = profile_service.import_raw_resume_file(
        db=db,
        filename="rx_resume.json",
        content_bytes=RX_RESUME_ACTUAL_FIXTURE.encode("utf-8"),
        mime_type="application/json",
        profile_id=master_profile.id,
    )

    updated_profile = profile_service.apply_raw_import_to_profile(
        db=db,
        import_id=raw_import.id,
        profile_id=master_profile.id,
    )

    # Assert all 8 CandidateProfile columns in database
    assert updated_profile.full_name == "Dr. Elena Rostova"
    assert updated_profile.email == "elena.rostova@quantum-labs.io"
    assert updated_profile.phone == "+1 (415) 888-9900"
    assert updated_profile.location == "San Francisco, CA"
    assert updated_profile.headline == "Principal AI Research Scientist & Quantum Engineer"
    assert updated_profile.summary == "Lead AI research scientist specializing in quantum-enhanced representation learning and distributed LLM alignment."
    assert updated_profile.linkedin_url == "https://linkedin.com/in/elena-rostova"
    assert updated_profile.github_url == "https://github.com/erostova"
    assert updated_profile.portfolio_url == "https://elenarostova.ai"

    # Invariant: Must be unverified
    assert updated_profile.is_verified is False
    assert len(updated_profile.experiences) == 1
    assert len(updated_profile.educations) == 1
    assert len(updated_profile.skills) >= 2


def test_parse_json_resume():
    parser = ResumeParserService()
    json_resume = """
    {
      "basics": {
        "name": "Sarah Connor",
        "email": "sarah@cyberdyne.org",
        "phone": "+1 555 999 0000",
        "headline": "Lead Defense Engineer",
        "summary": "Expert in cybernetic security systems."
      },
      "work": [
        {
          "company": "Cyberdyne Systems",
          "position": "Senior Engineer",
          "startDate": "2021-01",
          "endDate": "2023-12",
          "highlights": ["Designed defense core.", "Led 10 engineers."]
        }
      ],
      "education": [
        {
          "institution": "Caltech",
          "studyType": "B.S.",
          "area": "Computer Science"
        }
      ],
      "skills": [
        {"name": "Python", "level": "expert"},
        {"name": "C++", "level": "advanced"}
      ]
    }
    """
    result = parser.parse_raw_text(json_resume)
    assert result["provenance"] == "untrusted_import"
    assert result["is_verified"] is False
    assert result["profile"]["full_name"] == "Sarah Connor"
    assert result["profile"]["email"] == "sarah@cyberdyne.org"
    assert len(result["experiences"]) == 1
    assert result["experiences"][0]["company"] == "Cyberdyne Systems"
    assert result["experiences"][0]["is_verified"] is False
    assert len(result["skills"]) == 2
    assert result["skills"][0]["name"] == "Python"


def test_parse_docx_file():
    """Test format-aware binary DOCX extraction."""
    parser = ResumeParserService()
    
    # Create valid DOCX binary in memory
    doc = docx.Document()
    doc.add_heading("Marcus Aurelius", level=0)
    doc.add_paragraph("marcus.aurelius@rome.gov | +1 555 123 4567 | Rome, Empire")
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph("Experienced philosopher and executive leader specializing in stoic crisis management.")
    
    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph("Emperor - Roman Empire (161 – 180)")
    p1 = doc.add_paragraph(style="List Bullet")
    p1.add_run("Maintained Pax Romana during Marcomannic Wars.")
    p2 = doc.add_paragraph(style="List Bullet")
    p2.add_run("Authored Meditations on stoic resilience.")

    doc.add_heading("Education", level=1)
    doc.add_paragraph("Stoic Academy of Athens, Degree in Philosophy")

    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph("Strategy, Governance, Rhetoric, Stoicism, Diplomacy")

    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_bytes = docx_stream.getvalue()

    extracted_text, result = parser.parse_file_bytes(docx_bytes, filename="marcus_resume.docx", mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    
    # Invariant: Never extract PK\x03\x04 as candidate name
    assert "PK\x03\x04" not in result["profile"].get("full_name", "")
    assert result["profile"]["full_name"] == "Marcus Aurelius"
    assert result["profile"]["email"] == "marcus.aurelius@rome.gov"
    assert result["provenance"] == "untrusted_import"
    assert result["is_verified"] is False
    assert len(result["experiences"]) >= 1
    assert len(result["skills"]) >= 3


def test_parse_pdf_file():
    """Test format-aware PDF extraction."""
    parser = ResumeParserService()

    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=612, height=792)
    
    pdf_stream = io.BytesIO()
    writer.write(pdf_stream)
    blank_pdf_bytes = pdf_stream.getvalue()

    with pytest.raises(BadRequestError) as exc_info:
        parser.parse_file_bytes(blank_pdf_bytes, filename="blank.pdf", mime_type="application/pdf")
    assert "contained no extractable text" in str(exc_info.value)


def test_parse_malformed_docx_raises_bad_request():
    """Corrupted/invalid DOCX files must raise BadRequestError without corrupting profiles."""
    parser = ResumeParserService()
    corrupt_bytes = b"NOT_A_REAL_DOCX_FILE_GARBAGE"

    with pytest.raises(BadRequestError) as exc_info:
        parser.parse_file_bytes(corrupt_bytes, filename="fake.docx")
    assert "Malformed DOCX file" in str(exc_info.value)


def test_parse_malformed_pdf_raises_bad_request():
    """Corrupted/invalid PDF files must raise BadRequestError."""
    parser = ResumeParserService()
    corrupt_bytes = b"CORRUPTED_PDF_HEADER_12345"

    with pytest.raises(BadRequestError) as exc_info:
        parser.parse_file_bytes(corrupt_bytes, filename="corrupt.pdf")
    assert "Malformed PDF file" in str(exc_info.value)


def test_parse_plain_text_resume_never_invents_missing_facts():
    parser = ResumeParserService()
    text_resume = """
    Alexander Hamilton
    alexander@treasury.gov
    
    Summary
    Experienced financial architect and policy engineer.
    
    Skills
    Python, SQL, Financial Modeling, Distributed Ledgers
    
    Experience
    Senior Architect - First Bank of US (2020 - 2024)
    - Designed treasury infrastructure.
    - Automated ledger reconciliation.
    
    Education
    King's College, B.A. in Economics
    """
    result = parser.parse_raw_text(text_resume)
    assert result["provenance"] == "untrusted_import"
    assert result["is_verified"] is False
    assert result["profile"]["email"] == "alexander@treasury.gov"
    assert result["profile"]["phone"] is None
    assert len(result["skills"]) >= 3
    assert len(result["experiences"]) >= 1
    for exp in result["experiences"]:
        assert exp["is_verified"] is False


def test_rxresume_period_date_parsing_variations():
    """Verify that period strings in RxResume JSON format are accurately parsed without token loss or date swapping."""
    parser = ResumeParserService()

    payload = {
        "basics": {
            "name": "Marcus Aurelius",
            "email": "marcus@rome.org",
            "location": "Rome, Italy",
        },
        "sections": {
            "experience": {
                "items": [
                    {
                        "id": "exp-1",
                        "company": "Company A",
                        "position": "Role A",
                        "period": "March-2011 - January-2019",
                        "description": "Led operations."
                    },
                    {
                        "id": "exp-2",
                        "company": "Company B",
                        "position": "Role B",
                        "period": "March-2011 - Present",
                        "description": "Active leadership."
                    },
                    {
                        "id": "exp-3",
                        "company": "Company C",
                        "position": "Role C",
                        "period": "2011 - 2019",
                        "description": "Senior engineer."
                    },
                    {
                        "id": "exp-4",
                        "company": "Company D",
                        "position": "Role D",
                        "period": "March 2011 - January 2019",
                        "description": "Principal engineer."
                    },
                    {
                        "id": "exp-5",
                        "company": "Company E",
                        "position": "Role E",
                        "period": "2011",
                        "description": "Advisor."
                    },
                    {
                        "id": "exp-6",
                        "company": "Company F",
                        "position": "Role F",
                        "period": "2011-2019",
                        "description": "Staff engineer."
                    }
                ]
            },
            "education": {
                "items": [
                    {
                        "id": "edu-1",
                        "school": "University of Rome",
                        "degree": "B.S.",
                        "area": "Computer Science",
                        "period": "September-2007 - June-2011"
                    }
                ]
            }
        }
    }

    clean_text, result = parser.parse_file_bytes(
        json.dumps(payload).encode("utf-8"),
        filename="rx_dates.json",
        mime_type="application/json"
    )

    exps = result["experiences"]
    
    # 1. "March-2011 - January-2019"
    assert exps[0]["start_date"] == "March 2011"
    assert exps[0]["end_date"] == "January 2019"
    assert exps[0]["is_current"] is False

    # 2. "March-2011 - Present"
    assert exps[1]["start_date"] == "March 2011"
    assert exps[1]["end_date"] is None
    assert exps[1]["is_current"] is True

    # 3. "2011 - 2019"
    assert exps[2]["start_date"] == "2011"
    assert exps[2]["end_date"] == "2019"
    assert exps[2]["is_current"] is False

    # 4. "March 2011 - January 2019"
    assert exps[3]["start_date"] == "March 2011"
    assert exps[3]["end_date"] == "January 2019"
    assert exps[3]["is_current"] is False

    # 5. "2011"
    assert exps[4]["start_date"] == "2011"
    assert exps[4]["end_date"] is None

    # 6. "2011-2019"
    assert exps[5]["start_date"] == "2011"
    assert exps[5]["end_date"] == "2019"

    # 7. Education: "September-2007 - June-2011"
    edus = result["educations"]
    assert edus[0]["start_date"] == "September 2007"
    assert edus[0]["end_date"] == "June 2011"


def test_rxresume_period_rendered_accurately_in_compiler():
    """Verify that parsed dates render accurately in compiled HTML without date swapping or missing months."""
    from app.services.tailoring.compiler import ResumeDocumentCompiler
    compiler = ResumeDocumentCompiler()

    tailored_data = {
        "candidate": {
            "full_name": "Marcus Aurelius",
            "email": "marcus@rome.org",
        },
        "tailored_summary": {"text": "Experienced executive."},
        "highlighted_skills": [{"name": "Leadership"}],
        "tailored_experience": [
            {
                "company": "Roman Empire",
                "position": "Emperor",
                "start_date": "March 2011",
                "end_date": "January 2019",
                "is_current": False,
                "tailored_highlights": [{"text": "Maintained state stability."}],
            },
            {
                "company": "Senate Advisory",
                "position": "Chief Advisor",
                "start_date": "February 2019",
                "is_current": True,
                "tailored_highlights": [{"text": "Guiding strategy."}],
            }
        ]
    }

    compiled_html = compiler.compile_html(tailored_data=tailored_data, candidate_info=tailored_data["candidate"], educations=[])
    assert "March 2011 – January 2019" in compiled_html
    assert "February 2019 – Present" in compiled_html

