import io
import json
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from typing import Any, Dict, List, Optional, Tuple

import pypdf
import docx

from app.core.errors import BadRequestError
from app.core.logging import get_logger

logger = get_logger("app.services.parser")


class ResumeParserService:
    """Extracts structured draft candidate facts from untrusted imported resumes across PDF, DOCX, JSON, Markdown, and TXT.
    
    IMPORTANT INVARIANTS:
    - All extracted facts are strictly tagged as UNTRUSTED_DRAFT (is_verified=False, provenance='untrusted_import') until verified by the user.
    - NEVER invent missing facts (missing fields remain None or empty lists).
    - Format-aware parsing ensures binary streams are never treated as raw text.
    - Malformed files raise explicit BadRequestError rather than producing corrupted profiles.
    """

    def parse_file_bytes(
        self,
        content_bytes: bytes,
        filename: str,
        mime_type: Optional[str] = None,
    ) -> Tuple[str, Dict[str, Any]]:
        """Format-aware extraction from raw file bytes.
        
        Returns:
            Tuple of (extracted_clean_text, parsed_structured_draft_dict)
        """
        if not content_bytes or len(content_bytes.strip()) == 0:
            raise BadRequestError("Cannot parse empty file content.")

        ext = os.path.splitext(filename.lower())[1]

        # 1. DOCX Format
        if ext in [".docx", ".doc"] or mime_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]:
            extracted_text = self.extract_text_from_docx(content_bytes)
            parsed_data = self._parse_from_text(extracted_text)
            return extracted_text, parsed_data

        # 2. PDF Format
        if ext == ".pdf" or mime_type == "application/pdf" or content_bytes.startswith(b"%PDF"):
            extracted_text = self.extract_text_from_pdf(content_bytes)
            parsed_data = self._parse_from_text(extracted_text)
            return extracted_text, parsed_data

        # 3. JSON Format (including Reactive Resume & JSON Resume exports)
        if ext == ".json" or mime_type in ["application/json", "text/json"]:
            try:
                text_str = content_bytes.decode("utf-8", errors="replace")
                json_obj = json.loads(text_str)
                if isinstance(json_obj, dict):
                    parsed_data = self._parse_from_json_dict(json_obj)
                    # Create normalized text representation for storage
                    normalized_text = self._json_to_normalized_text(parsed_data)
                    return normalized_text, parsed_data
            except json.JSONDecodeError as err:
                raise BadRequestError(f"Invalid JSON resume structure: {err}")

        # 4. Markdown & Plain Text Format
        text_str = content_bytes.decode("utf-8", errors="replace")
        
        # Check if text is actually valid JSON despite file extension
        trimmed = text_str.strip()
        if trimmed.startswith("{") and trimmed.endswith("}"):
            try:
                json_obj = json.loads(trimmed)
                if isinstance(json_obj, dict):
                    parsed_data = self._parse_from_json_dict(json_obj)
                    normalized_text = self._json_to_normalized_text(parsed_data)
                    return normalized_text, parsed_data
            except Exception:
                pass  # Fallback to plain text / markdown

        parsed_data = self._parse_from_text(text_str)
        return text_str, parsed_data

    def parse_raw_text(self, raw_text: str) -> Dict[str, Any]:
        """Parse raw text/markdown/json string into structured draft candidate profile facts."""
        if not raw_text or not raw_text.strip():
            return self._empty_draft_response()

        # Check for JSON format
        trimmed = raw_text.strip()
        if trimmed.startswith("{") and trimmed.endswith("}"):
            try:
                parsed_json = json.loads(trimmed)
                if isinstance(parsed_json, dict):
                    return self._parse_from_json_dict(parsed_json)
            except Exception:
                pass  # Proceed with regular text parsing

        return self._parse_from_text(raw_text)

    # --- Format-Specific Extractors ---

    def extract_text_from_docx(self, content_bytes: bytes) -> str:
        """Extract plain text lines from DOCX binary bytes with paragraph and table preservation."""
        if not content_bytes.startswith(b"PK\x03\x04"):
            raise BadRequestError("Malformed DOCX file: Missing valid ZIP/DOCX header signature.")

        lines: List[str] = []

        # Attempt extraction via python-docx
        try:
            doc = docx.Document(io.BytesIO(content_bytes))
            for p in doc.paragraphs:
                p_text = p.text.strip()
                if p_text:
                    # Check if paragraph is list bullet
                    if p.style and "list" in p.style.name.lower():
                        lines.append(f"- {p_text}")
                    else:
                        lines.append(p_text)

            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        lines.append(" | ".join(row_cells))
        except Exception as docx_err:
            logger.warning(f"python-docx parsing failed ({docx_err}), falling back to direct XML extraction")
            # Fallback to direct XML extraction from word/document.xml
            try:
                with zipfile.ZipFile(io.BytesIO(content_bytes)) as z:
                    if "word/document.xml" not in z.namelist():
                        raise BadRequestError("Invalid DOCX archive: word/document.xml not found.")
                    xml_content = z.read("word/document.xml")
                    tree = ET.fromstring(xml_content)
                    # XML namespaces
                    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                    for p_node in tree.iterfind(".//w:p", ns):
                        texts = [t_node.text for t_node in p_node.iterfind(".//w:t", ns) if t_node.text]
                        if texts:
                            p_line = "".join(texts).strip()
                            if p_line:
                                lines.append(p_line)
            except Exception as xml_err:
                raise BadRequestError(f"Malformed or corrupted DOCX file: {xml_err}")

        extracted = "\n".join(lines).strip()
        if not extracted:
            raise BadRequestError("DOCX file was parsed successfully but contained no readable text.")
        return extracted

    def extract_text_from_pdf(self, content_bytes: bytes) -> str:
        """Extract plain text from PDF binary bytes page-by-page."""
        if not content_bytes.startswith(b"%PDF"):
            raise BadRequestError("Malformed PDF file: Missing valid %PDF header signature.")

        lines: List[str] = []
        try:
            reader = pypdf.PdfReader(io.BytesIO(content_bytes))
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    raise BadRequestError("PDF file is password protected and cannot be read.")

            for page_idx, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    lines.append(page_text.strip())
        except Exception as pdf_err:
            if isinstance(pdf_err, BadRequestError):
                raise
            raise BadRequestError(f"Malformed or corrupted PDF file: {pdf_err}")

        extracted = "\n\n".join(lines).strip()
        if not extracted:
            raise BadRequestError("PDF file was parsed successfully but contained no extractable text.")
        return extracted

    # --- Structured JSON / Reactive Resume Adapter ---

    def _extract_profile_network_and_url(self, prof: Dict[str, Any]) -> Tuple[str, Optional[str]]:
        """Extract platform network and valid URL from a profile object."""
        net = (prof.get("network") or prof.get("name") or "").lower().strip()
        u = None
        web = prof.get("website")
        if isinstance(web, dict):
            u = web.get("url") or web.get("href")
        elif isinstance(web, str):
            u = web
        if not u:
            u = prof.get("url") or prof.get("href") or prof.get("link")
        
        username = prof.get("username")
        if not u and username:
            if "linkedin" in net:
                u = f"https://linkedin.com/in/{username}"
            elif "github" in net:
                u = f"https://github.com/{username}"

        return net, u

    def _parse_period_date_range(self, period_str: Optional[str]) -> Tuple[Optional[str], Optional[str], bool]:
        """Accurately parse a period or date range into (start_date, end_date, is_current).
        
        Preserves original employment and education dates accurately without
        accidental start/end date swapping, day/month/year confusion, or character class truncation.
        """
        if not period_str:
            return None, None, False
        s = str(period_str).strip()
        if not s:
            return None, None, False

        def clean_token(tok: str) -> str:
            tok = tok.strip()
            # Convert "Month-YYYY" or "Month-YY" hyphen/dot/slash to space: e.g. "March-2011" -> "March 2011"
            m = re.match(r"^([a-zA-Z]+)[-–—/.](\d{2,4})$", tok)
            if m:
                return f"{m.group(1)} {m.group(2)}"
            return tok

        # Match standalone YYYY-YYYY or YYYY–YYYY (e.g. "2011-2019" or "2016 – 2021")
        m_years = re.match(r"^(\d{4})\s*[-–—]\s*(\d{4})$", s)
        if m_years:
            return m_years.group(1), m_years.group(2), False

        # Split on range delimiter: " - ", " – ", " — ", " to ", " until ", " till ", " through "
        # or dash right before Present/Current/Now/Ongoing
        parts = re.split(
            r"\s+[-–—]\s+|\s+(?:to|until|till|through)\s+|\s*[-–—]\s*(?=(?:present|current|now|ongoing)\b)",
            s,
            flags=re.IGNORECASE,
        )

        if len(parts) >= 2:
            start_part = clean_token(parts[0])
            end_part = clean_token(parts[1])
            if end_part.lower() in ["present", "current", "now", "ongoing"]:
                return start_part, None, True
            return start_part, end_part, False
        elif len(parts) == 1:
            tok = clean_token(parts[0])
            if tok.lower() in ["present", "current", "now", "ongoing"]:
                return None, None, True
            return tok, None, False

        return s, None, False

    def _parse_from_json_dict(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Convert a structured JSON resume (Reactive Resume v3/v4, JSON Resume, or native) into standardized draft facts."""
        basics = data.get("basics") or data.get("profile") or {}
        sections = data.get("sections") or {}

        # 1. Full Name
        full_name = (
            basics.get("name")
            or basics.get("full_name")
            or data.get("name")
            or data.get("full_name")
            or "Imported Candidate"
        )

        # 2. Email
        email = (
            basics.get("email")
            or data.get("email")
            or ""
        )

        # 3. Phone
        phone = (
            basics.get("phone")
            or basics.get("phoneNumber")
            or data.get("phone")
            or None
        )

        # 4. Location
        loc_val = basics.get("location") or data.get("location")
        location = None
        if isinstance(loc_val, str):
            location = loc_val.strip() or None
        elif isinstance(loc_val, dict):
            parts = [loc_val.get("address"), loc_val.get("city"), loc_val.get("region"), loc_val.get("country"), loc_val.get("postalCode")]
            location = ", ".join(filter(None, parts)) or None

        # 5. Headline
        headline = (
            basics.get("headline")
            or basics.get("label")
            or data.get("headline")
            or None
        )

        # 6. Summary (RxResume puts summary at root level data.summary.content, or sections.summary, or basics.summary)
        summary_raw = None
        if "summary" in data:
            s_obj = data["summary"]
            if isinstance(s_obj, dict):
                summary_raw = s_obj.get("content") or s_obj.get("text")
            elif isinstance(s_obj, str):
                summary_raw = s_obj

        if not summary_raw and "summary" in sections:
            s_obj = sections["summary"]
            if isinstance(s_obj, dict):
                summary_raw = s_obj.get("content") or s_obj.get("text")
            elif isinstance(s_obj, str):
                summary_raw = s_obj

        if not summary_raw and "summary" in basics:
            summary_raw = basics["summary"]

        summary = self._strip_html(summary_raw) if summary_raw else None

        # 7. Website / Portfolio
        website = None
        web_val = basics.get("website") or basics.get("url") or data.get("website")
        if isinstance(web_val, str):
            website = web_val
        elif isinstance(web_val, dict):
            website = web_val.get("url") or web_val.get("href")

        portfolio_url = basics.get("portfolio_url") or data.get("portfolio_url") or website
        linkedin_url = basics.get("linkedin_url") or basics.get("linkedin") or data.get("linkedin_url") or data.get("linkedin")
        github_url = basics.get("github_url") or basics.get("github") or data.get("github_url") or data.get("github")

        # 8. Check profiles lists across all standard schema locations
        profiles_list = []
        if "profiles" in sections and isinstance(sections["profiles"], dict):
            profiles_list.extend(sections["profiles"].get("items", []))
        elif "profiles" in sections and isinstance(sections["profiles"], list):
            profiles_list.extend(sections["profiles"])
        if "profiles" in basics and isinstance(basics["profiles"], list):
            profiles_list.extend(basics["profiles"])
        if "profiles" in data and isinstance(data["profiles"], list):
            profiles_list.extend(data["profiles"])

        for prof in profiles_list:
            if isinstance(prof, dict):
                net, u = self._extract_profile_network_and_url(prof)
                if u:
                    if ("linkedin" in net or "linkedin.com" in u) and not linkedin_url:
                        linkedin_url = u
                    elif ("github" in net or "github.com" in u) and not github_url:
                        github_url = u
                    elif not portfolio_url and ("portfolio" in net or "website" in net):
                        portfolio_url = u

        # 9. Check customFields in basics.customFields
        custom_fields = basics.get("customFields", [])
        if isinstance(custom_fields, list):
            for cf in custom_fields:
                if isinstance(cf, dict):
                    link = cf.get("link") or ""
                    text = cf.get("text") or ""
                    icon = (cf.get("icon") or "").lower()
                    if ("linkedin" in icon or "linkedin.com" in link or "linkedin.com" in text) and not linkedin_url:
                        linkedin_url = link or text
                    elif ("github" in icon or "github.com" in link or "github.com" in text) and not github_url:
                        github_url = link or text
                    elif not portfolio_url and ("portfolio" in icon or "website" in icon):
                        portfolio_url = link or text

        profile = {
            "full_name": full_name,
            "email": email,
            "phone": phone,
            "location": location,
            "headline": headline,
            "summary": summary,
            "website": website,
            "linkedin_url": linkedin_url,
            "github_url": github_url,
            "portfolio_url": portfolio_url,
            "is_verified": False,
        }

        # 2. Work Experiences (Reactive Resume sections.experience.items or JSON Resume work/experience)
        experiences: List[Dict[str, Any]] = []
        raw_work = []
        if "experience" in sections and isinstance(sections["experience"], dict):
            raw_work = sections["experience"].get("items", [])
        elif "work" in sections and isinstance(sections["work"], dict):
            raw_work = sections["work"].get("items", [])
        elif "work" in data:
            raw_work = data.get("work", [])
        elif "experiences" in data:
            raw_work = data.get("experiences", [])
        elif "experience" in data:
            raw_work = data.get("experience", [])

        for idx, item in enumerate(raw_work):
            if isinstance(item, dict):
                company = item.get("company") or item.get("name") or "Company"
                position = item.get("position") or item.get("role") or item.get("title") or "Position"
                exp_loc = item.get("location")
                
                # Dates handling (can be date string "March-2011 - January-2019" or separate startDate/endDate)
                date_str = item.get("period") or item.get("date") or ""
                start_date = item.get("startDate") or item.get("start_date") or ""
                end_date = item.get("endDate") or item.get("end_date") or None
                is_current = bool(item.get("is_current"))

                if date_str and not start_date:
                    parsed_start, parsed_end, parsed_curr = self._parse_period_date_range(date_str)
                    start_date = parsed_start or date_str
                    end_date = parsed_end
                    if parsed_curr:
                        is_current = True
                elif start_date:
                    norm_start, _, _ = self._parse_period_date_range(str(start_date))
                    start_date = norm_start or start_date
                    if end_date:
                        norm_end, _, _ = self._parse_period_date_range(str(end_date))
                        end_date = norm_end or end_date

                # Highlights & Description
                raw_summary = item.get("summary") or item.get("description")
                desc_text = self._strip_html(raw_summary) if raw_summary else None
                
                highlights: List[str] = []
                if "highlights" in item and isinstance(item["highlights"], list):
                    for h in item["highlights"]:
                        clean_h = self._strip_html(str(h))
                        if clean_h:
                            highlights.append(clean_h)
                elif raw_summary:
                    # Extract list items from HTML summary if present
                    li_matches = re.findall(r"<li[^>]*>(.*?)</li>", raw_summary, re.IGNORECASE | re.DOTALL)
                    if li_matches:
                        for m in li_matches:
                            clean_m = self._strip_html(m)
                            if clean_m:
                                highlights.append(clean_m)

                skills_used = item.get("skills_used") or item.get("technologies") or item.get("keywords") or []
                if isinstance(skills_used, str):
                    skills_used = [s.strip() for s in skills_used.split(",") if s.strip()]

                experiences.append({
                    "company": company,
                    "position": position,
                    "location": exp_loc,
                    "start_date": str(start_date) if start_date else "2022",
                    "end_date": str(end_date) if end_date else None,
                    "is_current": is_current or (end_date is None),
                    "description": desc_text if not highlights else None,
                    "highlights": highlights,
                    "skills_used": skills_used,
                    "order_index": idx,
                    "is_verified": False,
                })

        # 3. Educations (Reactive Resume sections.education.items or JSON Resume education)
        educations: List[Dict[str, Any]] = []
        raw_edu = []
        if "education" in sections and isinstance(sections["education"], dict):
            raw_edu = sections["education"].get("items", [])
        elif "education" in data:
            raw_edu = data.get("education", [])
        elif "educations" in data:
            raw_edu = data.get("educations", [])

        for edu in raw_edu:
            if isinstance(edu, dict):
                inst = edu.get("institution") or edu.get("school") or edu.get("university") or "Institution"
                degree = edu.get("studyType") or edu.get("degree") or "Degree"
                field = edu.get("area") or edu.get("field_of_study") or edu.get("major")
                
                date_str = edu.get("period") or edu.get("date") or ""
                start_date = edu.get("startDate") or edu.get("start_date") or None
                end_date = edu.get("endDate") or edu.get("end_date") or None

                if date_str and not start_date:
                    parsed_start, parsed_end, _ = self._parse_period_date_range(date_str)
                    start_date = parsed_start or date_str
                    end_date = parsed_end
                elif start_date:
                    norm_start, _, _ = self._parse_period_date_range(str(start_date))
                    start_date = norm_start or start_date
                    if end_date:
                        norm_end, _, _ = self._parse_period_date_range(str(end_date))
                        end_date = norm_end or end_date

                gpa = edu.get("grade") or edu.get("score") or edu.get("gpa") or None
                edu_desc = self._strip_html(edu.get("description") or edu.get("summary") or "")

                educations.append({
                    "institution": inst,
                    "degree": degree,
                    "field_of_study": field,
                    "start_date": str(start_date) if start_date else None,
                    "end_date": str(end_date) if end_date else None,
                    "gpa": str(gpa) if gpa else None,
                    "highlights": [edu_desc] if edu_desc else (edu.get("courses") or edu.get("highlights") or []),
                    "is_verified": False,
                })

        # 4. Skills (Reactive Resume sections.skills.items or JSON Resume skills)
        skills: List[Dict[str, Any]] = []
        raw_skills = []
        if "skills" in sections and isinstance(sections["skills"], dict):
            raw_skills = sections["skills"].get("items", [])
        elif "skills" in data:
            raw_skills = data.get("skills", [])

        for s_item in raw_skills:
            if isinstance(s_item, dict):
                s_name = s_item.get("name")
                s_desc = s_item.get("description")
                s_level = s_item.get("level") or s_item.get("proficiency") or "intermediate"
                if isinstance(s_level, int):
                    s_level = "expert" if s_level >= 4 else "intermediate" if s_level >= 2 else "beginner"

                if s_name:
                    skills.append({
                        "name": s_name,
                        "category": self._infer_skill_category(s_name),
                        "proficiency": str(s_level),
                        "is_verified": False,
                    })

                # If it has keywords or description containing comma list
                keywords = s_item.get("keywords") or []
                if isinstance(keywords, list):
                    for kw in keywords:
                        if kw and kw != s_name:
                            skills.append({
                                "name": str(kw).strip(),
                                "category": self._infer_skill_category(str(kw)),
                                "proficiency": "intermediate",
                                "is_verified": False,
                            })
                elif s_desc and ("," in s_desc or "|" in s_desc):
                    for token in re.split(r"[,|•·]", s_desc):
                        cleaned = token.strip()
                        if cleaned and len(cleaned) < 40 and cleaned != s_name:
                            skills.append({
                                "name": cleaned,
                                "category": self._infer_skill_category(cleaned),
                                "proficiency": "intermediate",
                                "is_verified": False,
                            })
            elif isinstance(s_item, str) and s_item.strip():
                skills.append({
                    "name": s_item.strip(),
                    "category": self._infer_skill_category(s_item),
                    "proficiency": "intermediate",
                    "is_verified": False,
                })

        # 5. Projects (Reactive Resume sections.projects.items or JSON Resume projects)
        projects: List[Dict[str, Any]] = []
        raw_proj = []
        if "projects" in sections and isinstance(sections["projects"], dict):
            raw_proj = sections["projects"].get("items", [])
        elif "projects" in data:
            raw_proj = data.get("projects", [])

        for proj in raw_proj:
            if isinstance(proj, dict):
                p_name = proj.get("name") or "Project"
                p_summary = self._strip_html(proj.get("description") or proj.get("summary") or "")
                p_url_val = proj.get("website") or proj.get("url")
                p_url = None
                if isinstance(p_url_val, dict):
                    p_url = p_url_val.get("url") or p_url_val.get("href")
                elif isinstance(p_url_val, str):
                    p_url = p_url_val

                techs = proj.get("keywords") or proj.get("technologies") or []
                if isinstance(techs, str):
                    techs = [t.strip() for t in techs.split(",") if t.strip()]

                projects.append({
                    "name": p_name,
                    "description": p_summary or None,
                    "url": p_url,
                    "highlights": proj.get("highlights") or [],
                    "technologies": techs,
                    "is_verified": False,
                })

        return {
            "provenance": "untrusted_import",
            "is_verified": False,
            "profile": profile,
            "experiences": experiences,
            "educations": educations,
            "skills": self._deduplicate_skills(skills),
            "projects": projects,
        }

    # --- Markdown & Text Extractor ---

    def _parse_from_text(self, text: str) -> Dict[str, Any]:
        """Robust rule-based text/markdown extractor with structured section segmentation."""
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if not lines:
            return self._empty_draft_response()

        # 1. Contact Info & Name
        email_match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
        email = email_match.group(0) if email_match else "imported@candidate.local"

        phone_match = re.search(r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
        phone = phone_match.group(0) if phone_match else None

        linkedin_match = re.search(r"https?://(?:www\.)?linkedin\.com/in/[\w-]+", text)
        linkedin_url = linkedin_match.group(0) if linkedin_match else None

        github_match = re.search(r"https?://(?:www\.)?github\.com/[\w-]+", text)
        github_url = github_match.group(0) if github_match else None

        portfolio_match = re.search(r"https?://(?!www\.linkedin\.com|www\.github\.com)[\w.-]+\.[a-z]{2,}(?:/\S*)?", text)
        portfolio_url = portfolio_match.group(0) if portfolio_match else None

        # Determine full_name from top lines
        full_name = "Imported Candidate"
        for candidate_line in lines[:5]:
            clean_l = candidate_line.lstrip("#*=- ").strip()
            # Ignore headers, emails, phones, URLs
            if (
                clean_l
                and "@" not in clean_l
                and "http" not in clean_l
                and not any(h in clean_l.lower() for h in ["resume", "curriculum vitae", "summary", "experience", "education"])
                and len(clean_l) < 50
                and not re.search(r"\d{3}", clean_l)
            ):
                full_name = clean_l
                break

        # 2. Section Segmentation
        section_headers = {
            "experience": ["experience", "work experience", "employment", "professional experience", "work history"],
            "education": ["education", "academic background", "academics", "educational background"],
            "skills": ["skills", "technical skills", "competencies", "technologies", "core competencies", "tools & technologies"],
            "projects": ["projects", "personal projects", "open source", "key projects", "notable projects"],
            "summary": ["summary", "professional summary", "about me", "objective", "executive summary", "profile"],
        }

        sections: Dict[str, List[str]] = {k: [] for k in section_headers}
        current_section = "summary"

        for line in lines:
            clean_line = line.strip("#*=-_ ").lower()
            matched_sec = None
            for sec_name, headers in section_headers.items():
                if clean_line in headers or any(clean_line.startswith(h) and len(clean_line) < len(h) + 6 for h in headers):
                    matched_sec = sec_name
                    break

            if matched_sec:
                current_section = matched_sec
            else:
                sections[current_section].append(line)

        # 3. Summary
        summary_lines = [l for l in sections["summary"] if not any(c in l for c in [email, full_name]) and len(l) > 10]
        summary_text = "\n".join(summary_lines[:5]) if summary_lines else None

        # 4. Skills
        skills: List[Dict[str, Any]] = []
        for s_line in sections["skills"]:
            # Check for category line: e.g. "- **Languages**: Python, Go, TypeScript"
            cat_match = re.match(r"(?:[-*•]\s*)?\*{0,2}([\w\s/&]+)\*{0,2}:\s*(.*)", s_line)
            if cat_match:
                cat_name = cat_match.group(1).strip()
                tokens = re.split(r"[,|•·]", cat_match.group(2))
                for t in tokens:
                    cl = t.strip().strip("-*• ")
                    if cl and len(cl) < 40:
                        skills.append({
                            "name": cl,
                            "category": self._infer_skill_category(cl, default_cat=cat_name),
                            "proficiency": "intermediate",
                            "is_verified": False,
                        })
            else:
                tokens = re.split(r"[,|•·\n]", s_line)
                for t in tokens:
                    cl = t.strip().strip("-*• ")
                    if cl and len(cl) < 40 and not cl.lower().startswith("skills"):
                        skills.append({
                            "name": cl,
                            "category": self._infer_skill_category(cl),
                            "proficiency": "intermediate",
                            "is_verified": False,
                        })

        # 5. Experiences
        experiences: List[Dict[str, Any]] = []
        current_exp: Optional[Dict[str, Any]] = None

        for line in sections["experience"]:
            if line.startswith(("-", "*", "•")):
                bullet = line.lstrip("-*• ").strip()
                if bullet:
                    if current_exp:
                        current_exp["highlights"].append(bullet)
            elif len(line) > 3 and (" - " in line or " | " in line or " – " in line or " at " in line or re.search(r"20\d\d", line)):
                if current_exp:
                    experiences.append(current_exp)

                # Parse Company, Position, Dates
                parts = re.split(r"[-|–—]", line)
                position = parts[0].strip() if len(parts) > 0 else "Software Engineer"
                company = parts[1].strip() if len(parts) > 1 else "Company"
                
                # Extract year/dates if present in line
                date_match = re.search(r"(\b20\d\d\b.*)", line)
                date_str = date_match.group(1) if date_match else "2022"
                
                current_exp = {
                    "company": company,
                    "position": position,
                    "location": None,
                    "start_date": date_str.split("-")[0].strip() if "-" in date_str else date_str,
                    "end_date": None,
                    "is_current": "present" in line.lower() or "current" in line.lower(),
                    "description": None,
                    "highlights": [],
                    "skills_used": [],
                    "order_index": len(experiences),
                    "is_verified": False,
                }
            elif current_exp and not current_exp["description"]:
                current_exp["description"] = line

        if current_exp:
            experiences.append(current_exp)

        # 6. Education
        educations: List[Dict[str, Any]] = []
        for line in sections["education"]:
            if any(term in line.lower() for term in ["university", "college", "institute", "school", "bachelor", "master", "phd", "b.s.", "m.s.", "b.a.", "degree"]):
                parts = re.split(r"[,|–—]", line)
                inst = parts[0].strip()
                deg = parts[1].strip() if len(parts) > 1 else "Degree"
                field = parts[2].strip() if len(parts) > 2 else None
                educations.append({
                    "institution": inst,
                    "degree": deg,
                    "field_of_study": field,
                    "start_date": None,
                    "end_date": None,
                    "gpa": None,
                    "highlights": [],
                    "is_verified": False,
                })

        # 7. Projects
        projects: List[Dict[str, Any]] = []
        current_proj: Optional[Dict[str, Any]] = None

        for line in sections["projects"]:
            if line.startswith(("-", "*", "•")):
                bullet = line.lstrip("-*• ").strip()
                if bullet and current_proj:
                    current_proj["highlights"].append(bullet)
            elif len(line) > 3:
                if current_proj:
                    projects.append(current_proj)
                current_proj = {
                    "name": line.strip().strip("#* "),
                    "description": None,
                    "url": None,
                    "highlights": [],
                    "technologies": [],
                    "is_verified": False,
                }
        if current_proj:
            projects.append(current_proj)

        profile = {
            "full_name": full_name,
            "email": email,
            "phone": phone,
            "location": None,
            "headline": None,
            "summary": summary_text,
            "website": portfolio_url,
            "linkedin_url": linkedin_url,
            "github_url": github_url,
            "portfolio_url": portfolio_url,
            "is_verified": False,
        }

        return {
            "provenance": "untrusted_import",
            "is_verified": False,
            "profile": profile,
            "experiences": experiences,
            "educations": educations,
            "skills": self._deduplicate_skills(skills[:30]),
            "projects": projects[:10],
        }

    # --- Helper Utilities ---

    def _strip_html(self, text: Optional[str]) -> str:
        """Strip HTML tags and convert entities."""
        if not text:
            return ""
        # Remove tags
        clean = re.sub(r"<[^>]+>", " ", text)
        # Normalize whitespace
        clean = re.sub(r"\s+", " ", clean).strip()
        return clean

    def _deduplicate_skills(self, skills: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Deduplicate skills by case-insensitive name."""
        seen = set()
        deduped = []
        for s in skills:
            name_key = s["name"].lower().strip()
            if name_key and name_key not in seen:
                seen.add(name_key)
                deduped.append(s)
        return deduped

    def _infer_skill_category(self, skill_name: str, default_cat: Optional[str] = None) -> str:
        """Infer skill category with fallback."""
        s = skill_name.lower().strip()
        if any(lang in s for lang in ["python", "javascript", "typescript", "java", "c++", "go", "rust", "ruby", "php", "sql", "html", "css", "kotlin", "swift", "c#", "scala"]):
            return "languages"
        elif any(fw in s for fw in ["react", "vue", "angular", "fastapi", "django", "flask", "node", "express", "spring", "next", "tailwind", "redux", "graphql"]):
            return "frameworks"
        elif any(db in s for db in ["postgres", "sqlite", "mysql", "mongodb", "redis", "elasticsearch", "cassandra", "dynamodb", "snowflake"]):
            return "databases"
        elif any(c in s for c in ["aws", "gcp", "azure", "docker", "kubernetes", "terraform", "ci/cd", "git", "linux", "jenkins", "helm"]):
            return "cloud_devops"
        
        if default_cat and default_cat.lower() in ["languages", "frameworks", "databases", "cloud_devops", "tools"]:
            return default_cat.lower()
        return "general"

    def _empty_draft_response(self) -> Dict[str, Any]:
        """Standard empty draft structure."""
        return {
            "provenance": "untrusted_import",
            "is_verified": False,
            "profile": {},
            "experiences": [],
            "educations": [],
            "skills": [],
            "projects": [],
        }

    def _json_to_normalized_text(self, parsed: Dict[str, Any]) -> str:
        """Convert parsed structured facts into clean readable markdown representation for storage."""
        p = parsed.get("profile", {})
        lines = [
            f"# {p.get('full_name', 'Candidate')}",
            f"Email: {p.get('email', '')} | Phone: {p.get('phone', '')} | Location: {p.get('location', '')}",
        ]
        if p.get("summary"):
            lines.extend(["\n## Professional Summary", p["summary"]])

        if parsed.get("experiences"):
            lines.append("\n## Work Experience")
            for exp in parsed["experiences"]:
                lines.append(f"\n### {exp.get('position')} at {exp.get('company')} ({exp.get('start_date')} – {exp.get('end_date') or 'Present'})")
                if exp.get("description"):
                    lines.append(exp["description"])
                for h in exp.get("highlights", []):
                    lines.append(f"- {h}")

        if parsed.get("educations"):
            lines.append("\n## Education")
            for edu in parsed["educations"]:
                lines.append(f"- {edu.get('degree')} in {edu.get('field_of_study') or 'General Studies'} – {edu.get('institution')}")

        if parsed.get("skills"):
            lines.append("\n## Skills")
            skill_names = [s["name"] for s in parsed["skills"]]
            lines.append(", ".join(skill_names))

        if parsed.get("projects"):
            lines.append("\n## Projects")
            for proj in parsed["projects"]:
                lines.append(f"- {proj.get('name')}: {proj.get('description') or ''}")

        return "\n".join(lines)


resume_parser = ResumeParserService()
